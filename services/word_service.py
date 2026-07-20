from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os


OUTPUT_FOLDER = "outputs"



def create_word(korean_text, english_text):


    os.makedirs(
        OUTPUT_FOLDER,
        exist_ok=True
    )


    document = Document()



    ####################################################
    # 제목
    ####################################################

    title = document.add_heading(
        "AI Meeting Minutes",
        level=1
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER



    ####################################################
    # 작성일
    ####################################################

    date = document.add_paragraph()

    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    date.add_run(
        f"Created : {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )



    document.add_paragraph(
        "================================================"
    )



    ####################################################
    # 한국어 회의록
    ####################################################

    document.add_heading(
        "■ 한국어 회의록",
        level=2
    )


    # Dictionary 처리

    if isinstance(korean_text, dict):

        add_korean_structured_minutes(
            document,
            korean_text
        )

    else:

        document.add_paragraph(
            korean_text
        )



    ####################################################
    # 구분선
    ####################################################

    document.add_paragraph(
        "------------------------------------------------"
    )



    ####################################################
    # 영문 회의록
    ####################################################

    document.add_heading(
        "■ English Meeting Minutes",
        level=2
    )


    document.add_paragraph(
        english_text
    )



    ####################################################
    # 저장
    ####################################################

    filename = os.path.join(
        OUTPUT_FOLDER,
        "meeting_minutes.docx"
    )


    document.save(
        filename
    )


    return filename





def add_korean_structured_minutes(
        document,
        minutes
):


    ####################################################
    # 기본 정보
    ####################################################

    sections = [

        (
            "회의 제목",
            minutes.get(
                "meeting_title",
                ""
            )
        ),

        (
            "회의 목적",
            minutes.get(
                "meeting_purpose",
                ""
            )
        ),

        (
            "참석자",
            minutes.get(
                "participants",
                ""
            )
        )

    ]



    for title, content in sections:


        document.add_heading(
            title,
            level=3
        )


        document.add_paragraph(
            content
        )



    ####################################################
    # 논의 사항
    ####################################################

    document.add_heading(
        "논의 사항",
        level=3
    )


    discussions = minutes.get(
        "discussion",
        []
    )


    for item in discussions:

        document.add_paragraph(
            f"• {item}"
        )



    ####################################################
    # 결정 사항
    ####################################################

    document.add_heading(
        "결정 사항",
        level=3
    )


    decisions = minutes.get(
        "decisions",
        []
    )


    for item in decisions:

        document.add_paragraph(
            f"• {item}"
        )



    ####################################################
    # Action Item 표
    ####################################################

    document.add_heading(
        "Action Items",
        level=3
    )


    actions = minutes.get(
        "action_items",
        []
    )


    if actions:


        table = document.add_table(
            rows=1,
            cols=3
        )


        table.alignment = WD_TABLE_ALIGNMENT.CENTER


        table.style = "Table Grid"



        headers = [
            "담당자",
            "업무",
            "완료 예정일"
        ]


        for i, text in enumerate(headers):

            table.rows[0].cells[i].text = text



        for action in actions:


            row = table.add_row().cells


            row[0].text = action.get(
                "owner",
                ""
            )

            row[1].text = action.get(
                "task",
                ""
            )

            row[2].text = action.get(
                "due_date",
                ""
            )



    ####################################################
    # 다음 일정
    ####################################################

    document.add_heading(
        "다음 일정",
        level=3
    )


    document.add_paragraph(
        minutes.get(
            "next_schedule",
            ""
        )
    )