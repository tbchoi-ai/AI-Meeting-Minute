from docx import Document
import os


OUTPUT_FOLDER = "output"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def save_word(meeting_minutes):

    document = Document()

    document.add_heading(
        "AI Meeting Minutes",
        level=1
    )

    document.add_paragraph(meeting_minutes)

    file_path = os.path.join(
        OUTPUT_FOLDER,
        "meeting_minutes.docx"
    )

    document.save(file_path)

    return file_path

from reportlab.platypus import SimpleDocTemplate
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet


def save_pdf(meeting_minutes):

    file_path = os.path.join(
        OUTPUT_FOLDER,
        "meeting_minutes.pdf"
    )

    styles = getSampleStyleSheet()

    doc = SimpleDocTemplate(file_path)

    story = []

    story.append(
        Paragraph(meeting_minutes, styles["Normal"])
    )

    doc.build(story)

    return file_path